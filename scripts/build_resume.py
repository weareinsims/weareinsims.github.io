from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Page margins: 0.75 inch all sides ---
for section in doc.sections:
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)

# --- Default style ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1a, 0x20, 0x35)


def set_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'),  str(after))
    if line:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    old = pPr.find(qn('w:spacing'))
    if old is not None:
        pPr.remove(old)
    pPr.append(spacing)


def add_bottom_border(para, color='1d4ed8', size=6):
    """Add a bottom border to a paragraph (acts as a section rule)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    old = pPr.find(qn('w:pBdr'))
    if old is not None:
        pPr.remove(old)
    pPr.append(pBdr)


def section_header(text):
    p = doc.add_paragraph()
    set_spacing(p, before=120, after=40)
    add_bottom_border(p, color='1d4ed8', size=6)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
    return p


def bullet(text, indent_left=0.25, hanging=0.2, size=10):
    p = doc.add_paragraph(style='Normal')
    set_spacing(p, before=0, after=20)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'),    str(int(indent_left * 1440)))
    ind.set(qn('w:hanging'), str(int(hanging * 1440)))
    pPr.append(ind)
    # Bullet character as a run then tab then text
    r_bullet = p.add_run('\u2022  ')
    r_bullet.font.size = Pt(size)
    r_bullet.font.name = 'Calibri'
    r_bullet.font.color.rgb = RGBColor(0x1a, 0x20, 0x35)
    r_text = p.add_run(text)
    r_text.font.size = Pt(size)
    r_text.font.name = 'Calibri'
    r_text.font.color.rgb = RGBColor(0x1a, 0x20, 0x35)
    return p


def job_header(title_company, date_range):
    """Bold job title | right-aligned date on the same line via tab stop."""
    p = doc.add_paragraph(style='Normal')
    set_spacing(p, before=100, after=20)
    # Tab stop at right margin
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9360')  # full content width for 0.75" margins on Letter
    tabs.append(tab)
    pPr.append(tabs)
    r1 = p.add_run(title_company)
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.name = 'Calibri'
    r1.font.color.rgb = RGBColor(0x1a, 0x20, 0x35)
    r2 = p.add_run('\t' + date_range)
    r2.bold = False
    r2.font.size = Pt(10)
    r2.font.name = 'Calibri'
    r2.font.color.rgb = RGBColor(0x5a, 0x64, 0x80)
    return p


def plain(text, size=10.5, bold=False, color=None, before=0, after=30):
    p = doc.add_paragraph(style='Normal')
    set_spacing(p, before=before, after=after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    return p


# ===================================================================
# NAME
# ===================================================================
name_p = doc.add_paragraph()
set_spacing(name_p, before=0, after=20)
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = name_p.add_run('Khizar Khan')
r.bold = True
r.font.size = Pt(22)
r.font.name = 'Calibri'
r.font.color.rgb = RGBColor(0x1a, 0x20, 0x35)

# ===================================================================
# CONTACT LINE
# ===================================================================
contact_p = doc.add_paragraph()
set_spacing(contact_p, before=0, after=60)
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_text = 'khizarazakhan@gmail.com  |  +1 (437) 575-1079  |  linkedin.com/in/khizarkhan1999/  |  weareinsims.github.io/ba-portfolio'
r = contact_p.add_run(contact_text)
r.font.size = Pt(9.5)
r.font.name = 'Calibri'
r.font.color.rgb = RGBColor(0x5a, 0x64, 0x80)

# ===================================================================
# PROFESSIONAL SUMMARY
# ===================================================================
section_header('Professional Summary')
summary = doc.add_paragraph(style='Normal')
set_spacing(summary, before=40, after=40)
r = summary.add_run(
    'Business analyst with a background in IT operations and hands-on experience translating '
    'technical complexity into structured requirements, process documentation, and stakeholder-ready '
    'deliverables. Comfortable working across the full BA lifecycle from discovery through to '
    'implementation support. Strong grasp of ITSM environments, having worked in and analyzed service '
    'desk operations from both the practitioner and analyst side.'
)
r.font.size = Pt(10.5)
r.font.name = 'Calibri'
r.font.color.rgb = RGBColor(0x1a, 0x20, 0x35)

# ===================================================================
# SKILLS
# ===================================================================
section_header('Skills')

skills = [
    ('Requirements Elicitation', 'Stakeholder interviews, workshops, gap analysis, user story writing'),
    ('Process Documentation',    'AS-IS/TO-BE mapping, BPMN, process flows, standard operating procedures'),
    ('BA Deliverables',          'BRD, functional specifications, use cases, user stories (Gherkin), personas'),
    ('Analysis',                 'Data analysis, KPI reporting, dashboard design, Excel'),
    ('ITSM Platforms',           'ServiceNow, Jira Service Management, Freshservice'),
    ('Collaboration Tools',      'Confluence, Lucidchart, Microsoft 365, SharePoint'),
    ('Methodologies',            'Agile, Scrum, ITIL'),
    ('Other',                    'SQL (basic), Python (basic), PowerShell (basic)'),
]

for label, detail in skills:
    p = doc.add_paragraph(style='Normal')
    set_spacing(p, before=0, after=20)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'),    str(int(0.0 * 1440)))
    pPr.append(ind)
    r1 = p.add_run(label + ': ')
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.name = 'Calibri'
    r1.font.color.rgb = RGBColor(0x1a, 0x20, 0x35)
    r2 = p.add_run(detail)
    r2.bold = False
    r2.font.size = Pt(10)
    r2.font.name = 'Calibri'
    r2.font.color.rgb = RGBColor(0x1a, 0x20, 0x35)

# ===================================================================
# BA PROJECTS
# ===================================================================
section_header('BA Projects')

projects = [
    {
        'title': 'IT Service Desk Process Improvement',
        'bullets': [
            'Conducted gap analysis comparing current service desk operations against ITIL best practices across 6 process areas, scoring each gap by impact and effort.',
            'Produced a full Business Requirements Document covering incident management, change control, knowledge management, and onboarding workflows.',
            'Developed a future-state proposal with prioritized recommendations, KPI targets, and a phased implementation plan.',
        ],
    },
    {
        'title': 'IT Help Desk Analytics Dashboard',
        'bullets': [
            'Analyzed 1,847 simulated tickets across a 6-month period to identify operational trends and performance gaps.',
            'Built an interactive dashboard with KPI cards, trend charts, and breakdown tables covering volume, resolution time, SLA compliance, FCR, and CSAT.',
            'Produced 5 written recommendations with estimated impact, each directly tied to findings in the data.',
        ],
    },
    {
        'title': 'ServiceNow ITSM Implementation BRD',
        'bullets': [
            'Developed a full BRD for a ServiceNow ITSM Pro implementation based on discovery interviews with 5 stakeholder groups.',
            'Documented 40+ functional requirements across 6 modules: Incident Management, Service Requests, Change Management, Asset Management, Self-Service Portal, and Reporting.',
            'Included risk register, success criteria, constraints, assumptions, implementation timeline, and a 15-item service catalog appendix.',
        ],
    },
    {
        'title': 'IT Onboarding Workflow - User Stories',
        'bullets': [
            'Developed 4 user personas covering new hires, hiring managers, IT agents, and IT managers, each with situation, goals, and frustrations.',
            'Wrote 10 user stories in Agile format with full Gherkin acceptance criteria (Given/When/Then), covering the complete onboarding and offboarding lifecycle including edge cases.',
        ],
    },
]

for proj in projects:
    p = doc.add_paragraph(style='Normal')
    set_spacing(p, before=80, after=20)
    r = p.add_run(proj['title'])
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(0x1a, 0x20, 0x35)
    for b in proj['bullets']:
        bullet(b)

# ===================================================================
# PROFESSIONAL EXPERIENCE
# ===================================================================
section_header('Professional Experience')

jobs = [
    {
        'title': 'Endpoint Technology Specialist  -  Centennial College',
        'date':  'Sept 2025 - Feb 2026',
        'bullets': [
            'Gathered operational requirements from IT and academic department stakeholders to define endpoint deployment standards and service expectations.',
            'Documented and maintained technical SOPs for patch deployment, endpoint imaging, and access provisioning workflows across 240+ managed devices.',
            'Analyzed deployment failure patterns and recurring incidents to identify root causes and recommend process improvements, reducing deployment lead time by 35 percent.',
            'Generated patch compliance and deployment status reports for IT leadership to support audit readiness, security reviews, and SLA accountability.',
            'Administered ServiceNow tickets and managed 50+ weekly incidents and service requests, maintaining 97 percent SLA compliance and a CSAT of 4.8 out of 5.',
            'Collaborated with infrastructure, security, networking, and AV teams on cross-functional issues, acting as a coordination point for escalated requests.',
        ],
    },
    {
        'title': 'IT Support Desk  -  Canadian National Institute for the Blind Foundation',
        'date':  'Apr 2025 - Aug 2025',
        'bullets': [
            'Assessed accessibility requirements for end users relying on assistive technologies (NVDA, JAWS) and translated these into support and training workflows.',
            'Delivered structured one-on-one and group training, improving user adoption and reducing repeat incidents.',
        ],
    },
    {
        'title': 'Service Desk Technician  -  Business Intelligence Analytics Inc.',
        'date':  'Sept 2023 - Sept 2024',
        'bullets': [
            'Conducted root cause analysis on recurring incidents and contributed to documentation and reporting improvements that raised first call resolution to 98 percent.',
            'Supported Azure Active Directory administration and policy-driven provisioning, documenting access workflows and control gaps.',
        ],
    },
    {
        'title': 'IT Technician / eLearning CAL Developer  -  Ontario Power Generation',
        'date':  'Jan 2021 - Aug 2021',
        'bullets': [
            'Analyzed support ticket data to identify recurring issues and recommend preventive solutions to reduce incident volume.',
            'Assisted with system access, provisioning, and remote support documentation across multiple internal departments.',
        ],
    },
]

for job in jobs:
    job_header(job['title'], job['date'])
    for b in job['bullets']:
        bullet(b)

# ===================================================================
# EDUCATION
# ===================================================================
section_header('Education')
job_header('HBCom in Cybersecurity & Information Technology  -  Toronto Metropolitan University', 'Sept 2019 - Oct 2024')

# ===================================================================
# CERTIFICATIONS
# ===================================================================
section_header('Certifications')

certs = [
    'CompTIA Security+',
    'ISC2 Certified in Cybersecurity (CC)',
    'Microsoft Azure Fundamentals (AZ-900)',
    'Microsoft Security Operations Analyst Associate',
    'Career Essentials in Cybersecurity by Microsoft & LinkedIn',
]
for c in certs:
    bullet(c)

# ===================================================================
# SAVE
# ===================================================================
out_path = '/Users/khizarkhan/projects/Portfolio/Resume_BA_KhizarKhan.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
